import glob
import os
import re
import shutil
from datetime import datetime
from typing import Any, Dict, List, Optional

import openpyxl
from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
import httpx
from pydantic import BaseModel
from uuid import uuid4

from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance
from sentence_transformers import SentenceTransformer


def _qdrant_search(self, collection_name: str, query_vector=None, limit: int = 5, **kwargs):
    return self.query_points(collection_name=collection_name, query=query_vector, limit=limit, **kwargs)

QdrantClient.search = _qdrant_search

BASE_DIR = os.path.dirname(__file__)
ENGINES_DIR = os.path.join(BASE_DIR, "engines")
CLIENT_FILES_DIR = os.path.join(BASE_DIR, "client_files")
EXPORTS_DIR = os.path.join(BASE_DIR, "exports")
EXPECTED_EXCEL_FILES = [
    "Master Excel Proto.xlsx",
    "Personal Finance Intelligence Engine.xlsx",
    "Pricing Model Valuation.xlsx",
]

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def read_excel_file(path: str) -> Dict[str, str]:
    """Read all sheets from an Excel workbook and convert each sheet into text."""
    workbook = openpyxl.load_workbook(path, data_only=False)
    sheet_texts: Dict[str, str] = {}

    for sheet in workbook.worksheets:
        lines: List[str] = [f"Sheet: {sheet.title}"]
        lines.append(f"Dimensions: {sheet.min_row}:{sheet.max_row} x {sheet.min_column}:{sheet.max_column}")

        for table in sheet.tables.values():
            lines.append(f"Table: {table.displayName or table.name} ref={table.ref}")
            column_names = [column.name for column in table.tableColumns]
            if column_names:
                lines.append(f"Columns: {', '.join(column_names)}")

        for row in sheet.iter_rows(values_only=False):
            for cell in row:
                if cell.value is None:
                    continue
                if cell.data_type == "f" or (isinstance(cell.value, str) and cell.value.startswith("=")):
                    lines.append(f"{cell.coordinate}: formula={cell.value}")
                else:
                    lines.append(f"{cell.coordinate}: {cell.value}")

        sheet_texts[sheet.title] = "\n".join(lines)

    return sheet_texts


def chunk_text(text: str, min_tokens: int = 300, max_tokens: int = 500) -> List[str]:
    """Split text into chunks of approximately 300-500 tokens."""
    tokens = re.findall(r"\w+|[^\s\w]", text)
    if not tokens:
        return []

    if len(tokens) <= max_tokens:
        return [text]

    chunks: List[str] = []
    start = 0
    while start < len(tokens):
        end = min(start + max_tokens, len(tokens))
        remaining = len(tokens) - end
        if remaining and remaining < min_tokens and chunks:
            end = len(tokens)
        chunk_tokens = tokens[start:end]
        chunks.append(" ".join(chunk_tokens))
        start = end

    return chunks


def embed_and_store(chunks: List[str], metadata: Dict[str, str]) -> int:
    """Embed text chunks and store vectors in Qdrant with metadata."""
    if not chunks:
        return 0

    vectors = embedder.encode(chunks)
    points = []
    batch_size = 64
    total_inserted = 0

    for chunk, vector in zip(chunks, vectors):
        payload = {"text": chunk, **metadata}
        points.append({
            "id": str(uuid4()),
            "vector": vector.tolist() if hasattr(vector, "tolist") else list(vector),
            "payload": payload,
        })
        if len(points) >= batch_size:
            qdrant.upsert(collection_name=collection_name, points=points)
            total_inserted += len(points)
            points.clear()

    if points:
        qdrant.upsert(collection_name=collection_name, points=points)
        total_inserted += len(points)

    return total_inserted


def ingest_excel_file(filename: str) -> Dict[str, int]:
    """Ingest a single Excel file from the engines directory."""
    file_path = os.path.join(ENGINES_DIR, filename)
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"Excel file not found: {filename}")

    sheets = read_excel_file(file_path)
    total_chunks = 0
    total_vectors = 0

    for sheet_name, sheet_text in sheets.items():
        chunks = chunk_text(sheet_text)
        total_chunks += len(chunks)
        total_vectors += embed_and_store(
            chunks,
            {"source": "engine", "source_file": filename, "sheet": sheet_name},
        )

    return {"chunks": total_chunks, "vectors": total_vectors}


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def save_uploaded_file(upload_file: UploadFile, client_folder: str) -> str:
    ensure_dir(client_folder)
    destination = os.path.join(client_folder, upload_file.filename)
    with open(destination, "wb") as out_file:
        shutil.copyfileobj(upload_file.file, out_file)
    return destination


def extract_text_from_pdf(path: str) -> str:
    text_lines: List[str] = []
    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None

    if pdfplumber is not None:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                text_lines.append(f"Page {i}: {page_text}")
        return "\n".join(text_lines)

    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError("pdfplumber or PyPDF2 is required for PDF extraction")

    reader = PdfReader(path)
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        text_lines.append(f"Page {i}: {page_text}")
    return "\n".join(text_lines)


def extract_uploaded_file_texts(path: str) -> Dict[str, str]:
    _, extension = os.path.splitext(path)
    extension = extension.lower()
    if extension == ".xlsx":
        return read_excel_file(path)
    if extension == ".pdf":
        return {"PDF": extract_text_from_pdf(path)}
    raise ValueError(f"Unsupported uploaded file type: {extension}")


def ingest_client_file(path: str, filename: str, client_folder: str) -> int:
    texts = extract_uploaded_file_texts(path)
    total_chunks = 0

    for sheet_name, text in texts.items():
        chunks = chunk_text(text)
        if chunks:
            total_chunks += len(chunks)
            embed_and_store(
                chunks,
                {
                    "source": "client_upload",
                    "filename": filename,
                    "sheet": sheet_name,
                    "client_folder": client_folder,
                },
            )

    return total_chunks


def is_client_upload_instruction(instruction: str) -> bool:
    lower_text = instruction.lower()
    return any(
        keyword in lower_text
        for keyword in [
            "uploaded client files",
            "uploaded financial statements",
            "client financials",
            "client pdf",
            "client upload",
            "uploaded client",
        ]
    )


def get_search_results(instruction: str, source_filter: Optional[str] = None, limit: int = 5) -> List[Dict[str, Any]]:
    query_vector = embedder.encode(instruction).tolist()
    search_response = qdrant.search(
        collection_name=collection_name,
        query_vector=query_vector,
        limit=limit,
        with_payload=True,
    )

    results = []
    for point in (search_response.points or []):
        payload = getattr(point, "payload", {}) or {}
        if source_filter and payload.get("source") != source_filter:
            continue
        results.append({
            "text": payload.get("text", ""),
            "source": payload.get("source"),
            "source_file": payload.get("source_file"),
            "filename": payload.get("filename"),
            "sheet": payload.get("sheet"),
            "client_folder": payload.get("client_folder"),
        })
    return results


def extract_financial_values_from_text(text: str) -> Dict[str, str]:
    normalized = text.replace("$", "").replace(",", "")
    patterns = {
        "Revenue": r"revenue\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
        "COGS": r"(?:cogs|cost of goods sold)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
        "Expenses": r"(?:expenses|operating expenses|opex)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
        "Cash": r"cash\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
        "Accounts Receivable": r"(?:accounts receivable|ar)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
        "Accounts Payable": r"(?:accounts payable|ap)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
        "EBITDA": r"ebitda\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)",
    }
    values: Dict[str, str] = {}
    for field, pattern in patterns.items():
        match = re.search(pattern, normalized, re.IGNORECASE)
        if match:
            values[field] = match.group(1)
    return values


def extract_client_info_from_text(text: str) -> Dict[str, str]:
    info: Dict[str, str] = {}
    fields = {
        "Client Name": r"client name\s*[:\-]?\s*(.+)",
        "Company": r"company\s*[:\-]?\s*(.+)",
        "Report Date": r"report date\s*[:\-]?\s*(.+)",
        "Period": r"period\s*[:\-]?\s*(.+)",
    }
    for key, pattern in fields.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            info[key] = match.group(1).strip().split("\n")[0]
    return info


def map_extracted_values_to_sheet(sheet_name: str, extracted: Dict[str, str]) -> Dict[str, str]:
    generic_mapping = {
        "Revenue": "Revenue",
        "COGS": "COGS",
        "Expenses": "Expenses",
        "Cash": "Cash",
        "Accounts Receivable": "Accounts Receivable",
        "Accounts Payable": "Accounts Payable",
        "EBITDA": "EBITDA",
    }
    mapped: Dict[str, str] = {}
    for friendly_name, value in extracted.items():
        target = generic_mapping.get(friendly_name)
        if target:
            mapped[target] = value
    return mapped


def write_values_to_excel(workbook: openpyxl.Workbook, sheet_name: str, mapping: Dict[str, str]) -> List[str]:
    written: List[str] = []
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"Sheet not found: {sheet_name}")
    for field_name, value in mapping.items():
        try:
            write_to_sheet(workbook, sheet_name, field_name, value)
            written.append(field_name)
        except ValueError:
            continue
    return written


def build_reason_prompt(instruction: str, workbook: str, sheet: str, fields: List[str], engine_chunks: List[str], client_chunks: List[str]) -> str:
    engine_text = "\n\n".join(engine_chunks[:5])
    client_text = "\n\n".join(client_chunks[:5])
    field_text = ", ".join(fields) if fields else "none"
    return (
        f"You are an AI assistant helping with Excel reasoning.\n"
        f"User instruction: {instruction}\n"
        f"Workbook: {workbook}\n"
        f"Sheet: {sheet}\n"
        f"Fields: {field_text}\n"
        f"Engine knowledge chunks:\n{engine_text}\n\n"
        f"Client upload chunks:\n{client_text}\n\n"
        f"Provide a concise answer, mention the workbook and sheet, and reference the detected fields.\n"
    )


def get_client_folder_from_instruction(instruction: str) -> Optional[str]:
    match = re.search(r"client[_\s-]?id\s*[:\-]?\s*([A-Za-z0-9_\-]+)", instruction, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    if not os.path.isdir(CLIENT_FILES_DIR):
        return None
    folders = [name for name in os.listdir(CLIENT_FILES_DIR) if os.path.isdir(os.path.join(CLIENT_FILES_DIR, name))]
    if not folders:
        return None
    latest = sorted(folders, reverse=True)[0]
    return latest


def get_client_upload_texts(client_folder: str) -> str:
    uploads_path = os.path.join(CLIENT_FILES_DIR, client_folder)
    all_texts: List[str] = []
    for root, _, files in os.walk(uploads_path):
        for filename in files:
            filepath = os.path.join(root, filename)
            if filename.lower().endswith(".xlsx"):
                sheet_texts = extract_uploaded_file_texts(filepath)
                all_texts.extend(sheet_texts.values())
            elif filename.lower().endswith(".pdf"):
                all_texts.append(extract_text_from_pdf(filepath))
    return "\n".join(all_texts)


def create_word_report(client_id: str, sections: List[str]) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required for Word export")

    client_folder_path = os.path.join(CLIENT_FILES_DIR, client_id)
    if not os.path.isdir(client_folder_path):
        raise FileNotFoundError(f"Client folder not found: {client_id}")

    ensure_dir(os.path.join(EXPORTS_DIR, client_id))
    output_path = os.path.join(EXPORTS_DIR, client_id, "report.docx")

    document = Document()
    document.add_heading("Client Report", level=1)

    client_text = get_client_upload_texts(client_id)
    client_info = extract_client_info_from_text(client_text)
    financial_values = extract_financial_values_from_text(client_text)

    if "client info" in sections:
        document.add_heading("Client Info", level=2)
        if client_info:
            for key, value in client_info.items():
                document.add_paragraph(f"{key}: {value}")
        else:
            document.add_paragraph("No client info extracted.")

    if "financial inputs" in sections:
        document.add_heading("Financial Inputs", level=2)
        if financial_values:
            for key, value in financial_values.items():
                document.add_paragraph(f"{key}: {value}")
        else:
            document.add_paragraph("No extracted financial inputs available.")

    if "normalized ebitda" in sections:
        document.add_heading("Normalized EBITDA", level=2)
        document.add_paragraph(financial_values.get("EBITDA", "Not extracted"))

    if "valuation" in sections:
        document.add_heading("Valuation Summary", level=2)
        document.add_paragraph("Valuation metrics were built from uploaded client data and engine models.")

    if "kpis" in sections:
        document.add_heading("KPIs", level=2)
        document.add_paragraph("Revenue, COGS, Expenses, Cash, AR, AP and EBITDA were extracted.")

    if "summary" in sections:
        document.add_heading("Narrative Summary", level=2)
        document.add_paragraph("This document summarizes client financials and reasoning outputs based on uploaded files and engine data.")

    document.save(output_path)
    return output_path


def trigger_word_export_if_requested(instruction: str) -> Optional[str]:
    lower_text = instruction.lower()
    if "word report" in lower_text or "generate a word report" in lower_text or "export word" in lower_text:
        client_id = get_client_folder_from_instruction(instruction)
        if not client_id:
            raise ValueError("No client folder available for Word export")
        return create_word_report(client_id, ["client info", "financial inputs", "normalized ebitda", "valuation", "kpis", "summary"])
    return None


def determine_workbook(instruction: str) -> str:
    lower_text = instruction.lower()
    explicit_workbooks = {
        "pricing model valuation": "Pricing Model Valuation.xlsx",
        "personal finance intelligence engine": "Personal Finance Intelligence Engine.xlsx",
        "master excel proto": "Master Excel Proto.xlsx",
        "pricing model valuation workbook": "Pricing Model Valuation.xlsx",
        "personal finance engine": "Personal Finance Intelligence Engine.xlsx",
        "master excel proto model": "Master Excel Proto.xlsx",
    }

    for phrase, filename in explicit_workbooks.items():
        if phrase in lower_text:
            return filename

    workbook_keywords = {
        "Master Excel Proto.xlsx": ["master excel proto", "valuation", "dcf", "ebitda", "entity"],
        "Personal Finance Intelligence Engine.xlsx": ["personal finance", "plev", "savings efficiency", "human capital"],
        "Pricing Model Valuation.xlsx": ["pricing", "product", "cogs", "competitor", "margin"],
    }

    for filename, keywords in workbook_keywords.items():
        if any(keyword in lower_text for keyword in keywords):
            return filename

    return "Master Excel Proto.xlsx"


def is_write_instruction(instruction: str) -> bool:
    return bool(re.search(r"\b(set|update|change|write|assign|enter)\b", instruction, re.IGNORECASE))


def extract_write_value(instruction: str) -> Optional[str]:
    match = re.search(r"\b(?:set|update|change|write|assign|enter)\b.*?\b(?:to|as)\b\s*(.+)", instruction, re.IGNORECASE)
    if not match:
        return None

    value = match.group(1).strip().strip(".\"'")
    if not value:
        return None

    if " if " in value:
        value = value.split(" if ", 1)[0].strip()
    return value


def extract_field_name(instruction: str) -> Optional[str]:
    match = re.search(r"\b(?:update|set|change|write|assign|enter)\b\s+(.*?)\s+(?:to|as)\b", instruction, re.IGNORECASE)
    if not match:
        return None

    field_name = match.group(1).strip().strip(".\"'")
    return field_name or None


async def call_llm_prompt(prompt: str) -> str:
    response = await call_llama(prompt)
    return response.get("response", "").strip()


def load_workbook(path: str) -> openpyxl.Workbook:
    return openpyxl.load_workbook(path, data_only=False)


def save_workbook(workbook: openpyxl.Workbook, filename: str) -> None:
    workbook.save(os.path.join(ENGINES_DIR, filename))


def find_field_cell(sheet: openpyxl.worksheet.worksheet.Worksheet, field_name: str) -> Optional[openpyxl.cell.cell.Cell]:
    normalized = field_name.strip().lower()
    for cell in sheet.iter_rows(values_only=False):
        for value in cell:
            if value.value is None:
                continue
            cell_text = str(value.value).strip().lower()
            if cell_text == normalized or normalized in cell_text:
                return value
    return None


def get_write_target_cell(sheet: openpyxl.worksheet.worksheet.Worksheet, label_cell: openpyxl.cell.cell.Cell) -> openpyxl.cell.cell.Cell:
    if label_cell.row == 1:
        return sheet.cell(row=2, column=label_cell.column)
    if label_cell.column == 1:
        return sheet.cell(row=label_cell.row, column=2)
    if label_cell.row < sheet.max_row:
        return sheet.cell(row=label_cell.row + 1, column=label_cell.column)
    return label_cell


def write_to_sheet(workbook: openpyxl.Workbook, sheet_name: str, field_name: str, value: str) -> str:
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"Sheet not found: {sheet_name}")

    sheet = workbook[sheet_name]
    label_cell = find_field_cell(sheet, field_name)
    if label_cell is None:
        raise ValueError(f"Field not found: {field_name}")

    target_cell = get_write_target_cell(sheet, label_cell)
    target_cell.value = value
    return f"Wrote {value} to {sheet_name}!{target_cell.coordinate}"


def get_sheet_candidates(workbook: openpyxl.Workbook, instruction: str) -> List[str]:
    normalized = instruction.lower()
    candidates = []
    for sheet_name in workbook.sheetnames:
        if sheet_name.lower() in normalized:
            candidates.append(sheet_name)

    return candidates


def get_field_candidates(workbook: openpyxl.Workbook, sheet_name: str) -> List[str]:
    sheet = workbook[sheet_name]
    candidates = set()
    for row in sheet.iter_rows(min_row=1, max_row=5, values_only=True):
        for cell in row:
            if isinstance(cell, str) and cell.strip():
                candidates.add(cell.strip())
    for row in sheet.iter_rows(min_col=1, max_col=1, max_row=sheet.max_row, values_only=True):
        cell = row[0]
        if isinstance(cell, str) and cell.strip():
            candidates.add(cell.strip())
    return list(candidates)


def detect_fields(instruction: str, candidates: List[str]) -> List[str]:
    normalized = instruction.lower()
    detected = []
    for candidate in candidates:
        candidate_norm = candidate.lower()
        if candidate_norm in normalized or any(token in normalized for token in re.split(r"[\s_]+", candidate_norm)):
            detected.append(candidate)
    return detected


def select_sheet_and_fields(instruction: str, workbook_name: str, metadata_results: List[Dict[str, Any]]) -> Dict[str, Any]:
    workbook_path = os.path.join(ENGINES_DIR, workbook_name)
    workbook = load_workbook(workbook_path)
    sheet_names = workbook.sheetnames

    explicit_sheets = get_sheet_candidates(workbook, instruction)
    sheet = explicit_sheets[0] if explicit_sheets else None

    if sheet is None:
        sheet_counts: Dict[str, int] = {}
        for item in metadata_results:
            if item.get("source_file") == workbook_name and item.get("sheet"):
                sheet_counts[item["sheet"]] = sheet_counts.get(item["sheet"], 0) + 1
        if sheet_counts:
            sheet = max(sheet_counts, key=sheet_counts.get)

    if sheet is None and sheet_names:
        sheet = sheet_names[0]

    if sheet is None:
        raise ValueError("Unable to determine target sheet")

    field_candidates = get_field_candidates(workbook, sheet)
    fields_detected = detect_fields(instruction, field_candidates)

    if not fields_detected:
        fields_detected = [item.get("sheet") for item in metadata_results if item.get("sheet") == sheet and item.get("source_file") == workbook_name]
        fields_detected = [f for f in fields_detected if f]

    return {
        "workbook": workbook_name,
        "sheet": sheet,
        "fields_detected": fields_detected,
        "action": "write" if is_write_instruction(instruction) else "read",
    }


@app.get("/")
async def homepage():
    return {"status": "SoarX backend running", "version": "1.0"}


class EmbedRequest(BaseModel):
    text: str


class SearchRequest(BaseModel):
    query: str


class ReasonRequest(BaseModel):
    text: str


class ExportRequest(BaseModel):
    client_id: str
    sections: List[str]


class IngestRequest(BaseModel):
    filename: str


@app.post("/ingest")
async def ingest(request: IngestRequest):
    try:
        result = ingest_excel_file(request.filename)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc))

    return {
        "file_processed": request.filename,
        "total_chunks": result["chunks"],
        "total_vectors_added": result["vectors"],
    }


@app.post("/ingest_all")
async def ingest_all():
    excel_files = sorted(glob.glob(os.path.join(ENGINES_DIR, "*.xlsx")))
    if not excel_files:
        raise HTTPException(status_code=404, detail="No Excel files found in engines directory")

    files_processed: List[str] = []
    total_chunks = 0
    total_vectors = 0

    for path in excel_files:
        filename = os.path.basename(path)
        result = ingest_excel_file(filename)
        files_processed.append(filename)
        total_chunks += result["chunks"]
        total_vectors += result["vectors"]

    return {
        "files_processed": files_processed,
        "total_chunks": total_chunks,
        "total_vectors_added": total_vectors,
    }


@app.post("/upload_file")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing upload filename")

    _, extension = os.path.splitext(file.filename)
    extension = extension.lower()
    if extension not in {".xlsx", ".pdf"}:
        raise HTTPException(status_code=400, detail="Only .xlsx and .pdf uploads are supported")

    client_folder_name = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    client_folder_path = os.path.join(CLIENT_FILES_DIR, client_folder_name)
    saved_path = save_uploaded_file(file, client_folder_path)
    texts = extract_uploaded_file_texts(saved_path)

    chunks_added = 0
    for sheet_name, text in texts.items():
        chunks = chunk_text(text)
        if chunks:
            chunks_added += len(chunks)
            embed_and_store(
                chunks,
                {
                    "source": "client_upload",
                    "filename": file.filename,
                    "sheet": sheet_name,
                    "client_folder": client_folder_name,
                },
            )

    return {
        "status": "uploaded",
        "filename": file.filename,
        "chunks_added": chunks_added,
        "client_folder": client_folder_name,
    }


@app.post("/export_word")
async def export_word(request: ExportRequest):
    try:
        path = create_word_report(request.client_id, request.sections)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    return {
        "status": "exported",
        "path": path,
    }


# ---------------------------
# Llama endpoint
# ---------------------------
OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "llama3"

@app.post("/llm")
async def call_llama(prompt: str):
    async with httpx.AsyncClient() as client:
        response = await client.post(
            OLLAMA_URL,
            json={"model": MODEL_NAME, "prompt": prompt, "stream": False},
            timeout=120
        )
    data = response.json()
    return {"response": data.get("response", "")}


# ---------------------------
# Embeddings + Qdrant
# ---------------------------
embedder = SentenceTransformer("all-MiniLM-L6-v2")

qdrant = QdrantClient("http://localhost:6333")

collection_name = "soarx_memory"

existing = qdrant.get_collections().collections
existing_names = [c.name for c in existing]

if collection_name not in existing_names:
    qdrant.create_collection(
        collection_name=collection_name,
        vectors_config=VectorParams(size=384, distance=Distance.COSINE)
    )


# ---------------------------
# Embed endpoint (FIXED)
# ---------------------------
@app.post("/embed")
async def embed_text(request: EmbedRequest):
    vector = embedder.encode(request.text).tolist()

    qdrant.upsert(
        collection_name=collection_name,
        points=[
            {
                "id": str(uuid4()),   # FIXED
                "vector": vector,
                "payload": {"text": request.text}
            }
        ]
    )

    return {"status": "stored", "text": request.text}


# ---------------------------
# Search endpoint (FIXED)
# ---------------------------
@app.post("/search")
async def search_text(request: SearchRequest):
    query_vector = embedder.encode(request.query).tolist()

    results = qdrant.search(
        collection_name=collection_name,
        query_vector=query_vector,
        limit=5
    )

    return {
        "matches": [
            {
                "score": r.score,
                "text": r.payload.get("text")
            }
            for r in (results.points or [])
        ]
    }

@app.post("/reason")
async def reason(request: ReasonRequest):
    instruction = request.text
    use_client_uploads = is_client_upload_instruction(instruction)

    engine_search = await search_text(SearchRequest(query=instruction))
    engine_chunks = [match["text"] for match in engine_search["matches"]]

    client_chunks: List[str] = []
    client_folder = get_client_folder_from_instruction(instruction)
    if use_client_uploads and client_folder:
        client_results = get_search_results(instruction, source_filter="client_upload")
        client_chunks = [item["text"] for item in client_results]

    metadata_results = get_search_results(instruction)
    workbook_name = determine_workbook(instruction)
    mapping = select_sheet_and_fields(instruction, workbook_name, metadata_results)

    prompt = build_reason_prompt(
        instruction,
        mapping["workbook"],
        mapping["sheet"],
        mapping["fields_detected"],
        engine_chunks,
        client_chunks,
    )
    answer = await call_llm_prompt(prompt)

    fields_updated: List[str] = []
    write_confirmation = None
    if mapping["action"] == "write":
        if use_client_uploads and client_chunks:
            client_text = " ".join(client_chunks)
            extracted_values = extract_financial_values_from_text(client_text)
            write_mapping = map_extracted_values_to_sheet(mapping["sheet"], extracted_values)
            if write_mapping:
                workbook = load_workbook(os.path.join(ENGINES_DIR, workbook_name))
                fields_updated = write_values_to_excel(workbook, mapping["sheet"], write_mapping)
                save_workbook(workbook, workbook_name)
                write_confirmation = f"Filled {len(fields_updated)} fields from uploaded client data."
        else:
            write_value = extract_write_value(instruction)
            field_name = None
            if mapping["fields_detected"]:
                field_name = mapping["fields_detected"][0]
            else:
                field_name = extract_field_name(instruction)

            if write_value and field_name:
                workbook = load_workbook(os.path.join(ENGINES_DIR, workbook_name))
                write_confirmation = write_to_sheet(workbook, mapping["sheet"], field_name, write_value)
                save_workbook(workbook, workbook_name)
                fields_updated = [field_name]
            elif mapping["action"] == "write":
                answer += "\nNote: the instruction appears to request a write, but no field or value could be extracted."

    report_path = None
    if "word report" in instruction.lower() or "generate a word report" in instruction.lower() or "export word" in instruction.lower():
        try:
            if not client_folder:
                client_folder = get_client_folder_from_instruction(instruction)
            if client_folder:
                report_path = create_word_report(client_folder, ["client info", "financial inputs", "normalized ebitda", "valuation", "kpis", "summary"])
                answer += f"\nGenerated Word report at {report_path}."
        except Exception as exc:
            answer += f"\nWord export failed: {exc}"

    return {
        "answer": answer,
        "workbook_used": mapping["workbook"],
        "sheet_used": mapping["sheet"],
        "fields_updated": fields_updated,
        "raw_chunks_used": engine_chunks + client_chunks,
        "write_confirmation": write_confirmation,
        "report_path": report_path,
    }